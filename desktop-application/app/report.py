import os
from multiprocessing import Pool
from functools import lru_cache
from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.shape import WD_INLINE_SHAPE
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import win32com.client
import pandas as pd
import pythoncom
import users
import questions
from tqdm import tqdm
import time

from PIL import Image, ImageFont, ImageDraw
import io

dial = Image.open("./desktop-application/app/graphics/dial.png")
pointer = Image.open("./desktop-application/app/graphics/pointer.png")
mf = ImageFont.truetype("./desktop-application/app/graphics/impact.ttf", 100)
wordchart = Image.open("./desktop-application/app/graphics/wordchart.png")
gradient = Image.open("./desktop-application/app/graphics/gradient/gradient.png")
tab = Image.open("./desktop-application/app/graphics/gradient/tab.png")

wordImportFile = pd.read_csv("./desktop-application/app/words.csv")
clusterImportFile = pd.read_csv("./desktop-application/app/clusters.csv")


wordList = []

class word:
    def __init__(self, name, ident):
        self.name = name
        self.ident = ident
        self.total = 0
        self.percent = 0

        #used for overall only
        self.stdd = 0
        self.stdp = 0

class cluster:
    def __init__(self, name, ident, words):
        self.name = name
        self.ident = ident
        self.words = words
        self.totalF = 0
        
        self.flag = False

@lru_cache(maxsize=1)
def get_template():
    template_path = os.path.join('desktop-application', 'app', 'report', 'template.docx')
    return Document(template_path)

def set_cell_color(cell, color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def convert_docx_to_pdf(docx_path, pdf_path):
    pythoncom.CoInitialize()
    word = win32com.client.Dispatch('Word.Application')
    try:
        doc = word.Documents.Open(docx_path)
        doc.SaveAs(pdf_path, FileFormat=17)  # FileFormat=17 is for PDF
        doc.Close()
    finally:
        word.Quit()
        pythoncom.CoUninitialize()

def initWords(wordList, wordImportFile):
    wordNameList = wordImportFile['word'].tolist()
    identList = wordImportFile['ident'].tolist()

    for i in wordImportFile.index:
        wordList.append(word(wordNameList[i], identList[i]))

def generateQuestionWeightedScore(userQuestions):
    cat = ["RFP", "EPS", "CM", "LdrSpv", "SrLdr"]
    scores = []
    overall = 0
    totalpossible = 0

    class CatScore:
        def __init__(self, name):
            self.name = name
            self.percentageScore = 0
            self.weightedScore = 0
            self.totalPossibleScore = 0
            self.dial = None

    for c in cat:
        scores.append(CatScore(c))

    for ques in questions.qList:
        for c in scores:
            if ques.qCat == c.name:
                if ques.qScore > 0:
                    c.totalPossibleScore += ques.qScore
                    totalpossible += ques.qScore

                if userQuestions[ques.qNum-1] == "Yes":
                    c.weightedScore += ques.qScore
                    overall += ques.qScore

                break
    for c in scores:
        c.percentageScore = int((c.weightedScore / c.totalPossibleScore)*100)

    overall = int((overall / totalpossible)*100)

    return scores, overall

def generateDial(dial, pointer, fnt, label, percent):
    width = 900
    ptext = str(int(percent))
    dial_copy = dial.copy()
    dialtext = ImageDraw.Draw(dial_copy)

    dialtext.text((width/ 2,630), label, (211,211,211), font=fnt, anchor="mm")
    dialtext.text((400,690), ptext, (0,0,0), font=fnt)

    rval = -1 * (-130 + (260 * (percent / 100)))

    pointerrotate = pointer.rotate(rval)
    intermediate = Image.alpha_composite(dial_copy, pointerrotate)
    return intermediate

def generateAllDials(dial, pointer, fnt, userQuestions):
    categories, overall = generateQuestionWeightedScore(userQuestions)
    dials = {}

    for cat in categories:
        name = cat.name
        pscore = cat.percentageScore
        cat.dial = generateDial(dial.copy(), pointer.copy(), fnt, name, pscore)
        dials[name] = cat.dial

    dials['Overall'] = generateDial(dial.copy(), pointer.copy(), fnt, "Overall", overall)
    return dials

def place_dials_on_document(doc, dials):
    # Find the paragraph with [Dials] and modify it
    target_paragraph = None
    for paragraph in doc.paragraphs:
        if '[Dials]' in paragraph.text:
            target_paragraph = paragraph
            break
    
    if target_paragraph is None:
        raise ValueError("Could not find [Dials] placeholder in the document.")

    # Remove the [Dials] placeholder
    target_paragraph.text = target_paragraph.text.replace('[Dials]', '')

    # Define the layout (left, top, width in inches)
    layout = [
        ('RFP', 2.9, 2, 2.25),
        ('Overall', 5.25, 1.6, 3),
        ('EPS', 8.35, 2, 2.25),
        ('CM', 3.25, 4.35, 2.25),
        ('SrLdr', 5.62, 4.65, 2.25),
        ('LdrSpv', 8, 4.35, 2.25)
    ]

    for (name, left, top, width) in layout:
        # Convert PIL Image to BytesIO object
        img_byte_arr = io.BytesIO()
        dials[name].save(img_byte_arr, format='PNG')
        img_byte_arr.seek(0)  # Move to the beginning of the BytesIO object

        # Add the image to the document
        run = target_paragraph.add_run()
        picture = run.add_picture(img_byte_arr, width=Inches(width))
        
        # Set the position of the image
        picture.left = Inches(left)
        picture.top = Inches(top)

    # Clear the dials dictionary to free up memory
    # Add page breaks after the dials

    #target_paragraph.add_run().add_break(WD_BREAK.PAGE)

    dials.clear()

def generateWordGraphic(arr, name, tUser, chart, fnt):
    newchart = chart.copy()
    newchart2 = chart.copy()

    seg1 = 0
    seg1T = 0
    seg1Words = []

    seg2 = 0
    seg2T = 0
    seg2Words = []

    seg3 = 0
    seg3T = 0
    seg3Words = []

    seg4 = 0
    seg4T = 0
    seg4Words = []

    pos = 0
    neg = 0
    for word in wordList:
        for w in arr:
            if w == word:
                if word.ident == "neg":
                    neg += 1
                    seg1 +=1

    for word in arr:
        if word.ident == "neg":
            neg += 1
            if (word.total / tUser) < .5:
                seg1 +=1
                seg1T += word.total
                seg1Words.append(word)
            else:
                seg2 +=1
                seg2T += word.total
                seg2Words.append(word)
        elif word.ident == "pos":
            pos += 1
            if (word.total / tUser) >= .5:
                seg3 +=1
                seg3T += word.total
                seg3Words.append(word)
            else:
                seg4 +=1
                seg4T += word.total
                seg4Words.append(word)




    #have totals for all words
    draw = ImageDraw.Draw(newchart)


    seg1p = str(int((seg1T / (seg1T+seg2T+seg3T+seg4T))*100)) + "%"
    seg2p = str(int((seg2T / (seg1T+seg2T+seg3T+seg4T))*100)) + "%"
    seg3p = str(int((seg3T / (seg1T+seg2T+seg3T+seg4T))*100)) + "%"
    seg4p = str(int((seg4T / (seg1T+seg2T+seg3T+seg4T))*100)) + "%"

    draw.text((200,650), "Words Selected:           " + str(seg1), "black", font=fnt)
    draw.text((200,250), "Words Selected:           " + str(seg2), "white", font=fnt)
    draw.text((1100,250), "Words Selected:           " + str(seg3), "white", font=fnt)
    draw.text((1100,650), "Words Selected:           " + str(seg4), "white", font=fnt)

    draw.text((200,750), "Percentage of Total:    " + seg1p, "black", font=fnt)
    draw.text((200,350), "Percentage of Total:    " + seg2p, "white", font=fnt)
    draw.text((1100,350), "Percentage of Total:    " + seg3p, "white", font=fnt)
    draw.text((1100,750), "Percentage of Total:    " + seg4p, "white", font=fnt)

    w = 1080

    ptext = name + "Word Association Summary"
    subtext = "n=" + str(tUser) + " participants"
    draw.text((w, 50), ptext, (0,0,0), font=fnt, anchor="mm")
    draw.text((w, 100), subtext, (0,0,0), font=fnt, anchor="mm")

    img_byte_arr = io.BytesIO()

    
    newchart.save(img_byte_arr, format='PNG')

    
    # Reset the BytesIO object positions
    img_byte_arr.seek(0)

    
    return img_byte_arr

def place_word_graphics_on_document(doc, user):
    # Find the paragraph with [WordMatrix] and modify it
    target_paragraph = None
    for paragraph in doc.paragraphs:
        if '[WordMatrix]' in paragraph.text:
            target_paragraph = paragraph
            break
    
    if target_paragraph is None:
        raise ValueError("Could not find [WordMatrix] placeholder in the document.")

    # Remove the [WordMatrix] placeholder
    target_paragraph.text = target_paragraph.text.replace('[WordMatrix]', '')

    # Generate the word graphics
    img_byte_arr = generateWordGraphic(user.words, "", 1 , wordchart, mf)

    # Add the images to the document side by side
    run = target_paragraph.add_run()
    
    picture1 = run.add_picture(img_byte_arr, width=Inches(3.5))


    # Set the position of the images (you may need to adjust these values)
    picture1.left = Inches(1)

def remove_empty_end_pages(doc):
    # Get all paragraphs and tables in the document
    all_elements = list(doc.paragraphs) + list(doc.tables)
    
    # Sort elements by their order in the document
    all_elements.sort(key=lambda x: x._element.getroottree().getpath(x._element))
    
    # Find the last non-empty element
    last_content_index = len(all_elements) - 1
    while last_content_index >= 0 and not element_has_content(all_elements[last_content_index]):
        last_content_index -= 1
    
    # Remove all elements after the last non-empty element
    for element in all_elements[last_content_index+1:]:
        parent = element._element.getparent()
        if parent is not None:
            parent.remove(element._element)

def element_has_content(element):
    if isinstance(element, Paragraph):
        return bool(element.text.strip())
    elif isinstance(element, Table):
        return any(cell.text.strip() for row in element.rows for cell in row.cells)
    return False

def initClusters(clusterList, words, clusterImportFile):
    clusterNameList = clusterImportFile['groupname'].tolist()
    clusterIdentList = clusterImportFile['ident'].tolist()
    wordList = clusterImportFile['words'].tolist()

    for i in range(len(clusterNameList)):
        clusterWords = [s.strip() for s in wordList[i].split(',')]
        temp = []
        for j in range(len(clusterWords)):
            tword = clusterWords[j]
            for wrd in words:
                if wrd.name == tword:
                    temp.append(word(wrd.name, wrd.ident))
                    break

        clusterList.append(cluster(clusterNameList[i], clusterIdentList[i], temp))

def processClusters(clusters, userwords):
    for words in userwords:
        for cls in clusters:
            for w in cls.words:
                if w.name == words:
                    cls.totalF += 1                   

def generate_individual_report(user):
    clusterList = []
    doc = get_template()
    initWords(wordList, wordImportFile)
    
    for paragraph in doc.paragraphs:
        if '[Name]' in paragraph.text:
            paragraph.text = paragraph.text.replace('[Name]', f"{user.fName} {user.lName}")

    # Generate the dials
    dials = generateAllDials(dial, pointer, mf, user.answers)
    
    # Place the dials on the document
    place_dials_on_document(doc, dials)

    # Place the word graphics on the document
    #place_word_graphics_on_document(doc, user)

    improvement_table = doc.tables[0]  # Assuming it's the first table
    
    ordered_categories = ['RFP', 'EPS', 'CM', 'LdrSpv', 'SrLdr']
    
    # Get subcategories from questions0
    subcategories = sorted(set(q.qSubCat for q in questions.qList))
    
    # Ensure the table has enough rows for all subcategories
    while len(improvement_table.rows) < len(subcategories) + 1:  # +1 for header row
        improvement_table.add_row()
    
    # Populate subcategories in the first column
    for i, subcat in enumerate(subcategories, start=1):
        improvement_table.cell(i, 0).text = subcat
    
     # Populate scores
    for i, subcat in enumerate(subcategories, start=1):
        for j, cat in enumerate(ordered_categories, start=1):
            cell = improvement_table.cell(i, j)
            score, total = calculate_score(user, questions.qList, subcat, cat)
            
            if score is not None and total > 0:
                percentage = (score / total) * 100
                #cell.text = f"{score} / {total}"
                color = get_color_hex(percentage)
            else:
                cell.text = ""
                color = "D3D3D3"  # Light gray for empty cells
            
            set_cell_color(cell, color)

    # Add a page break after the improvement table
    paragraph = improvement_table._element.getnext()
    if paragraph is None:
        paragraph = doc.add_paragraph()
    else:
        paragraph = Paragraph(paragraph, doc)
    run = paragraph.add_run()
    #run.add_break(WD_BREAK.PAGE)

    improvement_table = doc.tables[1]

    initClusters(clusterList, wordList, clusterImportFile)
    processClusters(clusterList, user.words)

    # Populate scores
    for i, clusters in enumerate(clusterList, start=1):
        cell = improvement_table.cell(i, 1)
        percentage = (clusters.totalF  / len(clusters.words)) * 100
        if clusters.ident == "n":
            percentage = 100 - percentage
        c = get_color_hex(percentage)
        set_cell_color(cell, c)


    #remove_empty_end_pages(doc)
    
    # Create the reports directory if it doesn't exist
    report_dir = os.path.join('desktop-application', 'app', 'report')
    os.makedirs(report_dir, exist_ok=True)
    
    temp_docx = os.path.abspath(os.path.join(report_dir, f"temp_{user.email}.docx"))
    doc.save(temp_docx)
    
    pdf_filename = os.path.abspath(os.path.join(report_dir, f"{user.email}.pdf"))
    convert_docx_to_pdf(temp_docx, pdf_filename)
    
    os.remove(temp_docx)
    
    return pdf_filename

def calculate_score(user, questions_list, subcat, cat):
    relevant_questions = [q for q in questions_list if q.qSubCat == subcat and q.qCat == cat]
    if not relevant_questions:
        return None, None
    
    yes_count = 0
    total_questions = len(relevant_questions)
    for question in relevant_questions:
        q_num = int(question.qNum)
        if q_num <= len(user.answers) and user.answers[q_num - 1].lower() == 'yes':
            yes_count += 1
    
    return yes_count, total_questions

def get_color_hex(percentage):
    if percentage <= 40:
        return "FF0000"  # Red
    elif 41 <= percentage <= 70:
        return "FFFF00"  # Yellow
    else:
        return "00FF00"  # Green

def generate_report_wrapper(user):
    try:
        return generate_individual_report(user)
    except Exception as e:
        print(f"Error generating report for {user.email}: {str(e)}")
        return None

def generate_reports_batch(users, batch_size=5):
    total_users = len(users)
    with tqdm(total=total_users, desc="Generating Reports", unit="report") as pbar:
        for i in range(0, total_users, batch_size):
            batch = users[i:i+batch_size]
            with Pool(os.cpu_count() - 1) as p:
                results = p.map(generate_report_wrapper, batch)
            
            # Update progress bar
            pbar.update(len(batch))
            
            # Filter out None results (failed report generations)
            successful_reports = [r for r in results if r is not None]
            #print(f"Successfully generated {len(successful_reports)} reports in this batch")
            time.sleep(1)

def generate_all_reports():
    users_with_scores = [user for user in users.userList if user.score != -1]
    generate_reports_batch(users_with_scores)
    #print(f"Attempted to generate reports for {len(users_with_scores)} users")

if __name__ == "__main__":
    generate_all_reports()